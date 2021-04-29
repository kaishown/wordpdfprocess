#enconding=utf-8
#author yanggx
#time 2021/4/19
#特殊符号映射查询
#https://blog.csdn.net/Hello_Mr_Zheng/article/details/99352176
#word read
#word save
from random import randint
from random import Random
from docx import Document
import random
import re
import os
import glob
from cn2capital import cncurrency

def load_itmes(file_name):
    with open(file_name,'r',encoding='utf-8') as fr:
        return fr.readlines()

class fill_docx:
    def __init__(self):
        max = 100
        self.items = {
            "card_id": self.random_multi_ids(max,18), # 身份证
            "contract_no":self.random_multi_CharInts(max, [2,16]),  # 编号
            "year":self.random_multi_i_j(max,2000, 2030),
            "month":self.random_multi_i_j(max,1, 12),
            "day":self.random_multi_i_j(max,1, 30),
            "amount":self.random_multi_float(max,1,10000),#金额
            "check_box":[chr(0xf0FE)],

        }
        self.items["amountCN"]= [cncurrency(str(i)) for i in self.items["amount"]]

    def random_with_N_digits(self,n):
        range_start = 10 ** (n - 1)
        range_end = (10 ** n) - 1
        return str(randint(range_start, range_end))
    def random_multi_ids(self,id_nums,n):
        ids=[]
        for i in range(id_nums):
            ids.append(self.random_with_N_digits(n))
        return ids
    def random_with_N_char(self,n):
        chars = ''
        english_chars = 'AaBbCcDdEeFfGgHhIiJjKkLlMmNnOoPpQqRrSsTtUuVvWwXxYyZz'
        length = len(english_chars) - 1
        random = Random()
        for i in range(n):
            chars += english_chars[random.randint(0, length)]
        return chars
    def random_multi_CharInts(self,id_nums,args):
        ids=[]
        for i in range(id_nums):
            chars = self.random_with_N_char(args[0])+self.random_with_N_digits(args[1])
            ids.append(chars)
        return ids
    def random_multi_i_j(self,id_nums,i,j):
        ids=[]
        for id in range(id_nums):
            ids.append(str(random.randint(i, j)))
        return ids
    def random_multi_float(self,id_nums,i,j):
        ids=[]
        for id in range(id_nums):
            ids.append(str(random.uniform(i,j)))
        return ids





def read_item_dir(dir):
    """
    读取目录下所有的文件
    :param dir:
    :return:
    """
    paths = os.path.join(dir,"*.txt")
    paths = glob.glob(paths)
    ITEMS={}
    for path in paths:
        item = os.path.split(path)[-1].strip(".txt")
        ITEMS[item] = load_itmes(path)
        random.shuffle(ITEMS[item])
    return ITEMS



def docx_pro(fill_item,document_file,save_path):
    document = Document(document_file)
    for indx, par in enumerate(document.paragraphs):
        # sub_texts = par.text.split("#")
        sub_texts = re.split("#",par.text)

        if "#" not in document.paragraphs[indx].text:
            continue
        else:
            document.paragraphs[indx].text = ''
            for sub_str in sub_texts:
                line=False
                font_name = False
                if sub_str and "@" in sub_str:
                    space = sub_str.split("@")[0]
                    sub_str=sub_str.split("@")[-1].strip()
                    if len(sub_str.split("/"))>1:
                        tag,i = sub_str.split("/")
                        print(tag)
                        sub_str = fill_item[tag][int(i)].strip()
                    else:
                        tag, i = sub_str.split("-")
                        line = True
                        # 勾选框添加字体
                        if tag == "check_box":
                            font_name = "Wingdings"
                            sub_str = space+fill_item[tag][int(i)-1]+space
                        else:
                            sub_str = space + fill_item[tag][int(i)].strip() + space
                run = document.paragraphs[indx].add_run(sub_str)
                # run.text = sub_str
                run.font.underline = line
                if font_name:
                    run.font.name = font_name


        print(document.paragraphs[indx].text)
    document.save(save_path)





# document_file='exu.docx'
def randomDoc(n,ori_doc):
    for i in range(n):
        items = read_item_dir("./items")
        other_items = fill_docx()
        fill_item = dict(items, **other_items.items)
        path = os.path.split(ori_doc)[0]
        save_path = os.path.join(path,str(i)+".docx")
        docx_pro(fill_item,ori_doc,save_path)


if __name__ == '__main__':
    randomDoc(1, 'exu.docx')





